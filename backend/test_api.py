"""
Script para verificar se a API e o conversor estão funcionando.
Gera um .docx mínimo, envia para POST /convert e valida a resposta.
"""
import io
import sys

from docx import Document

# Criar um .docx mínimo válido com uma tag
doc = Document()
doc.add_paragraph("#Introdução#")
doc.add_paragraph("Texto de teste da introdução.")
doc.add_paragraph("#Introdução#")

buffer = io.BytesIO()
doc.save(buffer)
buffer.seek(0)
docx_bytes = buffer.getvalue()

# Testar o conversor diretamente
try:
    from app.converter import convert
    result = convert(docx_bytes)
    assert isinstance(result, str), "convert deve retornar string"
    assert len(result) > 0, "resultado não pode ser vazio"
    assert "L-INTRODUCAO" in result or "Introdução" in result, "deve conter bloco de introdução"
    print("[OK] Conversor: convert() retornou texto válido.")
except Exception as e:
    print("[FALHA] Conversor:", e)
    sys.exit(1)

# Testar a API (requer servidor rodando ou vamos usar TestClient)
from fastapi.testclient import TestClient
from app.main import app

client = TestClient(app)

# GET /
r = client.get("/")
if r.status_code != 200:
    print("[FALHA] GET /:", r.status_code)
    sys.exit(1)
data = r.json()
if "message" not in data:
    print("[FALHA] GET /: resposta sem 'message'")
    sys.exit(1)
print("[OK] GET / retornou 200 e message.")

# POST /convert com .docx
r = client.post(
    "/convert",
    files={"file": ("teste.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
)
if r.status_code != 200:
    print("[FALHA] POST /convert:", r.status_code, r.text)
    sys.exit(1)
if not r.content:
    print("[FALHA] POST /convert: corpo vazio")
    sys.exit(1)
cd = r.headers.get("Content-Disposition")
if not cd or "attachment" not in cd:
    print("[AVISO] Content-Disposition não é attachment:", cd)
text = r.content.decode("utf-8")
if "L-INTRODUCAO" not in text and "Introdução" not in text:
    print("[AVISO] Resposta pode não conter bloco esperado; conteúdo:", text[:200])
print("[OK] POST /convert retornou 200 e arquivo .txt.")

# Validações de erro
r = client.post("/convert", files={"file": ("x.txt", b"xxx", "text/plain")})
if r.status_code != 400:
    print("[FALHA] Esperado 400 para .txt, obteve:", r.status_code)
else:
    print("[OK] Extensão inválida retorna 400.")

r = client.post("/convert", files={"file": ("vazio.docx", b"", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")})
if r.status_code != 400:
    print("[FALHA] Esperado 400 para arquivo vazio, obteve:", r.status_code)
else:
    print("[OK] Arquivo vazio retorna 400.")

print("\nTudo funcionando.")
