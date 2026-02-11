"""
Conversor DOCX para E-Book (texto com tags).
Entrada: caminho de arquivo ou bytes do DOCX.
Saída: string do conteúdo gerado (UTF-8).
"""

import io
import re
from pathlib import Path
from typing import List, Union

from docx import Document


class ConversionError(Exception):
    """Erro ao converter o documento."""

    pass


def convert(source: Union[str, Path, bytes]) -> str:
    """
    Converte um arquivo DOCX com tags no formato E-Book.

    Args:
        source: Caminho do arquivo .docx (str ou Path) ou bytes do documento.

    Returns:
        Conteúdo do arquivo de texto gerado (uma string, linhas separadas por \\n).

    Raises:
        ConversionError: Se o documento for inválido ou a conversão falhar.
    """
    if isinstance(source, bytes):
        doc = Document(io.BytesIO(source))
    else:
        doc = Document(source)

    html = _process_document(doc)
    return "\n".join(html)


def _process_document(doc: Document) -> list:
    paragrafosTotais = len(doc.paragraphs)
    if paragrafosTotais == 0:
        return []

    html = []
    proxParagrafo = 0
    boleano = True
    listaParagrafos = []
    listaAuxiliar = []
    paragrafoIndice = -1
    marcaNoTexto = []
    padrao = re.compile(r"^(.*?)\s")
    padraoFormula = re.compile(r"{{(.*?)}}")

    for paragrafo in doc.paragraphs:
        paragrafoStrip = paragrafo.text.strip()
        if "#" in paragrafoStrip and paragrafoStrip.startswith("%"):
            paragrafoStrip = f"#{paragrafoStrip.split('#')[1].split('#')[0]}#"

        if (
            paragrafoStrip.startswith("#")
            and paragrafoStrip.endswith("#")
            and proxParagrafo < paragrafosTotais - 1
            and marcaNoTexto == []
        ):
            marcaNoTexto.append(paragrafoStrip.lower())
            if not doc.paragraphs[proxParagrafo + 1].text.startswith("#"):
                listaParagrafos = []
                while boleano:
                    proxParagrafo += 1
                    if (
                        not doc.paragraphs[proxParagrafo].text.isspace()
                        and doc.paragraphs[proxParagrafo].text != ""
                    ):
                        listaParagrafos.append(proxParagrafo)
                        listaAuxiliar.append(proxParagrafo)
                    if proxParagrafo < paragrafosTotais - 1:
                        next_text = doc.paragraphs[proxParagrafo + 1].text
                        if not next_text.isspace() and next_text != "":
                            boleano = (
                                not next_text.strip().startswith("#")
                                and not next_text.strip().startswith("%")
                            )
                    else:
                        boleano = False

        boleano = True
        paragrafoIndice += 1
        proxParagrafo = paragrafoIndice + 1

        if paragrafoStrip.lower() in marcaNoTexto:
            marcaNoTexto.append(paragrafoStrip.lower())

        if paragrafoIndice in listaAuxiliar or len(marcaNoTexto) == 3:
            if paragrafoStrip.lower() in marcaNoTexto:
                marcaNoTexto = []

        result = ""
        try:
            result = _process_paragraph(
                doc,
                paragrafoStrip,
                paragrafoIndice,
                proxParagrafo,
                paragrafosTotais,
                listaParagrafos,
                listaAuxiliar,
                padrao,
                padraoFormula,
            )
        except Exception as e:
            raise ConversionError(f"Erro ao processar parágrafo: {e}") from e

        if isinstance(result, list):
            for linha in result:
                if linha != "" and (not html or linha != html[-1]):
                    html.append(linha)
        elif result != "" and result != "<p></p>" and (not html or result != html[-1]):
            html.append(result)

        listaParagrafos = []

    return html


def _process_paragraph(
    doc,
    paragrafoStrip,
    paragrafoIndice,
    proxParagrafo,
    paragrafosTotais,
    listaParagrafos,
    listaAuxiliar,
    padrao,
    padraoFormula,
) -> Union[str, List[str]]:
    texto = ""

    if paragrafoStrip.upper()[:7] == "UNIDADE":
        texto = f'<h2 class="title-vg">{paragrafoStrip.upper()}</h2>'

    elif paragrafoStrip and paragrafoStrip[0].isdigit():
        resultado = padrao.match(paragrafoStrip)
        if resultado:
            valorTitulo = resultado.group(1)
            if len(valorTitulo.replace(".", "")) == 1:
                texto = paragrafoStrip.replace(f"{valorTitulo} ", '<h4 class="subtitlei-vg">')
                texto = f"{texto}</h4>"
            elif len(valorTitulo.replace(".", "")) == 2:
                texto = paragrafoStrip.replace(f"{valorTitulo} ", "<p></p><h5 class=\"subtitleii-vg\">")
                texto = f"{texto}</h5>"
            else:
                texto = paragrafoStrip.replace(f"{valorTitulo} ", "<p></p><h6 class=\"subtitleiii-vg\">")
                texto = f"{texto}</h6>"

    elif "#introdução#" in paragrafoStrip.lower():
        texto = '<div class="L-INTRODUCAO" data-interaction="true"><div>{"texto":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '"}</div>Introdução</div>'

    elif "#conclusão#" in paragrafoStrip.lower():
        texto = '<div class="L-CONCLUSION" data-interaction="true"><div>{"texto":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '"}</div>Conclusão</div>'

    elif "#referências#" in paragrafoStrip.lower():
        texto = '<div class="L-REFERENCIASBB" data-interaction="true"><div>{"texto":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                texto = _formatar_texto_estilo(doc, indice, texto)
                if "<>" in texto:
                    link = doc.paragraphs[indice].text.split("<")[1].split(">")[0]
                    texto = texto.replace("<>", "%3Ca%20href='" + link + "'%20target='_blank'%20rel='noopener'%3E" + link + "%3C/a%3E")
            texto += '"}</div>Referências Bibliográficas</div><p></p>'

    elif "#apresentação#" in paragrafoStrip.lower() or "#destaque#" in paragrafoStrip.lower():
        texto = '<div class="D-DESTAQUE" data-interaction="true"><div>{"destaque":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '"}</div>Destaque</div>'

    elif "#citação#" in paragrafoStrip.lower():
        texto = '<div class="D-CDIRETA" data-interaction="true"><div>{"texto":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '"}</div>Recuo</div>'

    elif "#caixa#" in paragrafoStrip.lower():
        texto = '<div class="D-CAIXA" data-interaction="true"><div>{"caixa":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                if indice == listaParagrafos[0]:
                    texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A%3Cul%3E%0A"
                else:
                    texto += "%3Cli%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3Cbr%20/%3E%3Cbr%20/%3E%3C/li%3E%0A"
        texto += '%3C/ul%3E"}</div>Caixa</div>'

    elif "#glossário#" in paragrafoStrip.lower():
        texto = '<div class="U-SEARCHBLOCK" data-interaction="true"><div>{"title":"Glossário","conteudo":"%3Col%3E%0A'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                textoAuxiliar = textoAuxiliar.replace(": ", ": %3C/strong%3E")
                texto += "%3Cli%20style='text-align:%20justify;'%3E%3Cstrong%3E" + textoAuxiliar + "%3Cbr%20/%3E%3Cbr%20/%3E%3C/li%3E%0A"
        texto += '%3C/ol%3E"}</div>Bloco Busca</div>'

    elif "#video#" in paragrafoStrip.lower():
        conteudo = paragrafoStrip.replace("#Video#", "")
        texto = '<div class="T-VIDEO" data-interaction="true"><div>{{"link":"{}","video":"","pdf":""}}</div>Vídeo</div>'.format(conteudo)

    elif paragrafoStrip.lower().startswith("figura "):
        figura = "https://i.pinimg.com/736x/be/09/97/be0997e2d5732322bf552c6f2883c86e.jpg"
        titulo = doc.paragraphs[paragrafoIndice].text.split(":")
        fonte = doc.paragraphs[paragrafoIndice + 1].text.split(":")
        listaAuxiliar.append(proxParagrafo)
        texto = '<div class="T-FIGURA" data-interaction="true"><div>{{"titulo":"{}","fonte":"{}","accessibility":"{}","imagem":"{}"}}</div>Figura</div>'.format(
            titulo[1], fonte[1], titulo[1], figura
        )

    elif paragrafoStrip.lower().startswith("quadro "):
        tabela = "%3Ctable%20style='border-collapse:%20collapse;%20width:%20100%25;'%20border='1'%3E%0A%3Ctbody%3E%0A%3Ctr%3E%0A%3Ctd%20style='width:%20100%25;'%3EQuadro%3C/td%3E%0A%3C/tr%3E%0A%3C/tbody%3E%0A%3C/table%3E"
        titulo = doc.paragraphs[paragrafoIndice].text.split(":")
        fonte = doc.paragraphs[paragrafoIndice + 1].text.split(":")
        listaAuxiliar.append(proxParagrafo)
        texto = '<div class="T-QUADRO" data-interaction="true"><div>{{"titulo":"{}","fonte":"{}","texto":"{}","accessibility":"{}","quadro":""}}</div>Quadro</div>'.format(
            titulo[1], fonte[1], tabela, titulo[1]
        )

    elif "#reflita#" in paragrafoStrip.lower():
        texto = '<div class="B-REFLITA" data-interaction="true"><div>{"conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '","link":"","pdf":""}</div>Reflita</div>'

    elif "#saiba mais#" in paragrafoStrip.lower():
        texto = '<div class="U-SAIBAMAIS" data-interaction="true"><div>{"conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                if doc.paragraphs[indice].text.startswith("http"):
                    texto += '","link":"{}","pdf":""}}</div>Saiba Mais</div>'.format(doc.paragraphs[indice].text)
                else:
                    textoAuxiliar = _formatar_box_texto(doc, indice)
                    texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"

    elif "#atenção#" in paragrafoStrip.lower():
        texto = '<div class="U-ATENCAO" data-interaction="true"><div>{"conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '","pdf":""}</div>Atenção</div>'

    elif "#dica#" in paragrafoStrip.lower():
        texto = '<div class="B-DICA" data-interaction="true"><div>{"conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '","imagem":""}</div>Dica</div>'

    elif "#técnico#" in paragrafoStrip.lower():
        texto = '<div class="D-PROGRAMACAO" data-interaction="true"><div>{"texto":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '"}</div>Técnico</div>'

    elif (
        "#dica de livro#" in paragrafoStrip.lower()
        or "#dica do professor#" in paragrafoStrip.lower()
        or "#dica de leitura#" in paragrafoStrip.lower()
    ):
        if "#dica de livro#" in paragrafoStrip.lower():
            texto = '<div class="B-GREEN" data-interaction="true"><div>{"titulo":"Dica de Livro","conteudo":"'
        elif "#dica de leitura#" in paragrafoStrip.lower():
            texto = '<div class="B-GREEN" data-interaction="true"><div>{"titulo":"Dica de Leitura","conteudo":"'
        else:
            texto = '<div class="B-GREEN" data-interaction="true"><div>{"titulo":"Dica do Professor(a)","conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
        texto += '","imagem":"https://dbunicv.realize.pro.br/files/bbe7a7d8253ab4d19c641e74a008e50d.jpg"}</div>Esquerda</div>'

    elif (
        "#indicação de filme#" in paragrafoStrip.lower()
        or "#na web#" in paragrafoStrip.lower()
        or "#dica de filme#" in paragrafoStrip.lower()
    ):
        link = ""
        if "#indicação de filme#" in paragrafoStrip.lower() or "#dica de filme#" in paragrafoStrip.lower():
            texto = '<div class="B-BLUE" data-interaction="true"><div>{"titulo":"Indicação de Filme","conteudo":"'
        else:
            texto = '<div class="B-BLUE" data-interaction="true"><div>{"titulo":"Na Web","conteudo":"'
        if listaParagrafos:
            for indice in listaParagrafos:
                textoAuxiliar = doc.paragraphs[indice].text
                if "http" not in textoAuxiliar:
                    textoAuxiliar = _formatar_box_texto(doc, indice)
                    texto += "%3Cp%20style='text-align:%20justify;'%3E" + textoAuxiliar + "%3C/p%3E%0A"
                else:
                    link = textoAuxiliar
        texto += '","imagem":"https://dbunicv.realize.pro.br/files/851d95d42b89c5b0b24155447cf81d6b.jpg"}</div>Direita</div>'
        texto += '<div class="T-VIDEO" data-interaction="true"><div>{{"link":"{}","video":"","pdf":""}}</div>Vídeo</div>'.format(link)

    elif "#infográfico interativo#" in paragrafoStrip.lower():
        if listaParagrafos:
            texto = ""
            for indice in listaParagrafos:
                textoAuxiliar = _formatar_box_texto(doc, indice)
                elementos = textoAuxiliar.split(":", 1)
                texto += '<div class="I-ZSANFONA" data-interaction="true"><div>{"titulo":"<strong>' + elementos[0] + ':</strong>","conteudo":"%3Cp%20style=\'text-align:%20justify;\'%3E' + elementos[1] + '%3C/p%3E%0A"}</div>Sanfona</div>\n'
            texto = texto[: texto.rfind("\n")]

    elif "#forca" in paragrafoStrip.lower():
        conteudo = paragrafoStrip.replace("#Forca", "")
        texto = '<div class="I-JOGOFORCA" data-interaction="true"><div>{{"palavra":"{}"}}</div>Forca</div><p></p>'.format(conteudo)

    else:
        if proxParagrafo - 1 < 0:
            return ""
        p = doc.paragraphs[proxParagrafo - 1]
        texto = ""
        for run in p.runs:
            if (run.font.subscript or run.font.superscript) and run.text != " ":
                if run.font.subscript:
                    run.text = run.text.replace(run.text, f"<sub>{run.text}</sub>")
                else:
                    run.text = run.text.replace(run.text, f"<sup>{run.text}</sup>")
            else:
                if run.bold and run.text != " ":
                    run.text = run.text.replace(run.text, f"<strong>{run.text}</strong>")
                if run.italic and run.text != " ":
                    run.text = run.text.replace(run.text, f"<em>{run.text}</em>")
                if run.underline:
                    run.text = run.text.replace(run.text, f"<u>{run.text}</u>")
            texto += run.text

        if p.style.name.startswith("List"):
            texto = f'<li style="text-align: justify;">{texto}<br/><br/></li>'
        else:
            if not texto.isspace() and texto != "":
                texto = f'<p style="text-align: justify;">{texto}</p>'
            else:
                texto = f"<p>{texto}</p>"

        if "\n" in texto:
            texto = texto.replace("\n", "</p>|<p style=\"text-align: justify;\">")

        if "strong>" in texto:
            texto = texto.replace("<strong><strong>", "<strong>")
            texto = texto.replace("</strong></strong>", "</strong>")
            texto = texto.replace("</strong><strong>", "")
            texto = texto.replace("</strong> <strong>", " ")
        if "em>" in texto:
            texto = texto.replace("</strong></em><em><strong>", "")
            texto = texto.replace("<em><em>", "<em>")
            texto = texto.replace("</em></em>", "</em>")
            texto = texto.replace("</em> </em>", " ")
            texto = texto.replace("</em> <em>", " ")
            texto = texto.replace("</em><em>", "")

        if "{{" in texto:
            formulas = padraoFormula.findall(texto)
            for formula in formulas:
                formulaFormatada = _substituir_formula(formula)
                texto = texto.replace("{{" + formula + "}}", f"<em>{formulaFormatada}</em>")

        if "</p>|<p" in texto:
            partes = texto.split("|")
            out = []
            for linha in partes:
                if "<p style=\"text-align: justify;\"></p>" in linha:
                    linha = linha.replace("<p style=\"text-align: justify;\"></p>", "<p></p>")
                out.append(linha)
            return out
        if "<p></p>" in texto or "<p> </p>" in texto:
            pass
        else:
            return texto

    return texto


def _formatar_box_texto(doc, indice: int) -> str:
    texto_formatado = ""
    numero_aspas = 0
    for run in doc.paragraphs[indice].runs:
        if run.bold and run.italic and run.underline:
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(
                run.text,
                f"%3Cspan%20style=%22text-decoration:%20underline;%22%3E%3Cstrong%3E%3Cem%3E{run.text}%3C/em%3E%3C/strong%3E%3C/span%3E",
            )
        elif run.bold and run.underline:
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(
                run.text,
                f"%3Cspan%20style=%22text-decoration:%20underline;%22%3E%3Cstrong%3E{run.text}%3C/strong%3E%3C/span%3E",
            )
        elif run.bold and run.italic and run.text != " ":
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(run.text, f"%3Cstrong%3E%3Cem%3E{run.text}%3C/em%3E%3C/strong%3E")
        elif run.bold and run.text != " ":
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(run.text, f"%3Cstrong%3E{run.text}%3C/strong%3E")
        elif run.italic and run.text != " ":
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(run.text, f"%3Cem%3E{run.text}%3C/em%3E")
        elif run.underline:
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
            run.text = run.text.replace(
                run.text,
                f"%3Cspan%20style=%22text-decoration:%20underline;%22%3E{run.text}%3C/span%3E",
            )
        if ('"' in run.text or "%" in run.text) and not run.bold and not run.italic and not run.underline:
            if '"' in run.text:
                run.text, numero_aspas = _formatar_aspas(run.text, numero_aspas)
            if "%" in run.text:
                run.text = run.text.replace("%", "%25")
        if "%3C/span%3E%3Cspan%20style=%22text-decoration:%20underline;%22%3E" in run.text:
            run.text = run.text.replace(
                "%3C/span%3E%3Cspan%20style=%22text-decoration:%20underline;%22%3E",
                "",
            )
        texto_formatado += run.text
    return texto_formatado


def _formatar_aspas(texto: str, numero_aspas: int) -> tuple:
    """Corrigido: paridade com numero_aspas % 2 == 0 (não /2)."""
    if '"' in texto:
        if numero_aspas % 2 == 0:
            texto = texto.replace('"', '"')
            numero_aspas += 1
        else:
            texto = texto.replace('"', '"')
            numero_aspas -= 1
    return texto, numero_aspas


def _formatar_texto_estilo(doc, indice: int, texto: str) -> str:
    texto_formatado = ""
    for run in doc.paragraphs[indice].runs:
        if run.bold and run.text != " ":
            run.text = run.text.replace(run.text, f"%3Cstrong%3E{run.text}%3C/strong%3E")
        if run.italic and run.text != " ":
            run.text = run.text.replace(run.text, f"%3Cem%3E{run.text}%3C/em%3E")
        texto_formatado += run.text
    texto += "%3Cp%20style='text-align:%20justify;'%3E" + texto_formatado + "%3C/p%3E%0A"
    if "strong>" in texto:
        texto = texto.replace("%3Cstrong%3E%3Cstrong%3E", "%3Cstrong%3E")
        texto = texto.replace("%3C/strong%3E%3C/strong%3E", "%3C/strong%3E")
        texto = texto.replace("%3C/strong%3E%3Cstrong%3E", "")
        texto = texto.replace("%3C/strong%3E %3Cstrong%3E", " ")
    if "em>" in texto:
        texto = texto.replace("%3C/strong%3E%3C/em%3E%3Cem%3E%3Cstrong%3E", "")
        texto = texto.replace("%3Cem%3E%3Cem%3E", "%3Cem%3E")
        texto = texto.replace("%3C/em%3E%3C/em%3E", "%3C/em%3E")
        texto = texto.replace("%3C/em%3E %3C/em%3E", " ")
        texto = texto.replace("%3C/em%3E %3Cem%3E", " ")
        texto = texto.replace("%3C/em%3E%3Cem%3E", "")
    return texto


def _substituir_formula(formula: str) -> str:
    caracteres_especiais = ["+", "-", "=", ":", "*", "/", ")"]
    formula_formatada = ""
    sup = False
    sub = False
    union = False
    for caractere in formula.strip():
        if caractere == "^":
            sup = True
            if sub:
                sub = False
                caractere = "</sub><sup>"
            else:
                caractere = "<sup>"
        elif caractere == "_":
            sub = True
            if sup:
                sup = False
                caractere = "</sup><sub>"
            else:
                caractere = "<sub>"
        elif sup or sub:
            if caractere in caracteres_especiais:
                if sup and not union:
                    sup = False
                    caractere = "</sup>" + caractere
                elif sub and not union:
                    sub = False
                    caractere = "</sub>" + caractere
            if caractere == "(":
                union = True
            elif caractere == ")":
                union = False
                if sup:
                    sup = False
                    caractere += "</sup>"
                elif sub:
                    sub = False
                    caractere += "</sub>"
        formula_formatada += caractere
    formula_formatada = (
        formula_formatada.replace("<sub> </sub>", "").replace("<sup> </sup>", "").replace(" ", "")
    )
    return formula_formatada
